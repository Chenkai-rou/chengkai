import streamlit as st
# ... 其他 import 保持不变 ...

# 1. 注入自定义 CSS 皮肤 (放在最前面)
st.markdown("""
    <style>
    /* 整体背景与文字颜色 */
    .stApp {
        background-color: #0e1117;
        color: #ffffff;
    }
    /* 侧边栏样式 */
    [data-testid="stSidebar"] {
        background-image: linear-gradient(#2e333d, #0e1117);
        border-right: 1px solid #4f4f4f;
    }
    /* 输入框霓虹边框效果 */
    .stTextInput > div > div > input {
        border: 1px solid #00ffcc !important;
        box-shadow: 0 0 10px #00ffcc;
    }
    /* 按钮样式优化 */
    .stButton > button {
        background-color: #00ffcc !important;
        color: #000000 !important;
        border-radius: 20px;
        font-weight: bold;
        transition: 0.3s;
    }
    .stButton > button:hover {
        box-shadow: 0 0 20px #00ffcc;
        transform: scale(1.05);
    }
    </style>
    """, unsafe_allow_stdio=True, unsafe_allow_html=True)

# 2. 页面设置
st.set_page_config(page_title="Cyber Kai 5.0", page_icon="🌙", layout="wide")
# ... 后续逻辑保持不变 ...
from openai import OpenAI
from gtts import gTTS
from docx import Document
from io import BytesIO
# 新增的图像处理库
import easyocr
from PIL import Image
import numpy as np

# 1. 页面设置
st.set_page_config(page_title="Cyber Kai 4.0 OCR", page_icon="🏀", layout="wide")
st.title("🏀 程凯 | 智能识图版")

# --- 高级技巧：缓存 OCR 模型，避免重复加载 ---
@st.cache_resource
def load_ocr_model():
    # 加载中文和英文识别模型（第一次运行会比较慢）
    reader = easyocr.Reader(['ch_sim', 'en'], gpu=False)
    return reader

# 加载模型（这行代码在程序启动时会卡住几秒钟，正常的）
ocr_reader = load_ocr_model()
# -------------------------------------------

# 2. 自动获取密钥
try:
    api_key = st.secrets["DEEPSEEK_API_KEY"]
except:
    api_key = st.sidebar.text_input("请输入 DeepSeek API Key", type="password")

# 3. 功能函数（语音和导出保持不变）
def speak_text(text):
    try:
        tts = gTTS(text=text, lang='zh-cn')
        tts.save("response.mp3")
        audio_file = open("response.mp3", "rb")
        audio_bytes = audio_file.read()
        st.audio(audio_bytes, format="audio/mp3", autoplay=True)
    except Exception as e:
        st.warning(f"语音生成失败 (可能是权限问题，云端通常正常): {e}")

def export_to_word(chat_history):
    doc = Document()
    doc.add_heading('程凯聊天记录', 0)
    for msg in chat_history:
        role_name = "程凯" if msg["role"] == "assistant" else "我"
        doc.add_paragraph(f"{role_name}: {msg['content']}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 侧边栏：工具箱 (新增摄像头) ---
st.sidebar.title("🛠️ 多模态工具箱")

# === 新功能：拍照识别 ===
st.sidebar.write("---")
st.sidebar.header("📷 拍照识字")
# 调用摄像头组件
img_file_buffer = st.sidebar.camera_input("拍一张带有文字的照片")
ocr_result_text = ""

if img_file_buffer is not None:
    with st.spinner("正在努力识别图片中的文字..."):
        # 1. 读取图片
        image = Image.open(img_file_buffer)
        # 2. 转换为 numpy 数组供 OCR 使用
        img_array = np.array(image)
        # 3. 开始识别
        results = ocr_reader.readtext(img_array)
        # 4. 提取文字结果
        text_list = [res[1] for res in results]
        ocr_result_text = "\n".join(text_list)
        
        if ocr_result_text:
            st.sidebar.success("识别成功！")
            # 把识别出的文字显示出来，方便用户复制或直接发送
            st.sidebar.text_area("识别结果 (可复制)", ocr_result_text, height=100)
        else:
            st.sidebar.warning("如果你拍了照但没识别出来，可能是字太潦草或者光线太暗。")

# === 导出功能区 ===
st.sidebar.write("---")
st.sidebar.header("💾 数据管理")
if st.session_state.get("messages"):
    word_data = export_to_word(st.session_state.messages)
    st.sidebar.download_button(
        label="📥 导出聊天记录为 Word",
        data=word_data,
        file_name="chat_history.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if st.sidebar.button("🗑️ 清空当前对话"):
        st.session_state.messages = []
        st.rerun()
# ---------------------------------


# 6. 聊天记录显示
if "messages" not in st.session_state:
    st.session_state["messages"] = []
    st.session_state.messages.append({"role": "assistant", "content": "你好，我是程凯。除了聊天，现在我还装上了眼睛，可以帮你看看简单的文字。"})

for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

# 7. 交互逻辑
# 如果有 OCR 识别结果，提示用户是否要发送
initial_value = ""
if ocr_result_text:
    initial_value = f"我刚拍了一张照片，里面识别出的文字是：\n---\n{ocr_result_text}\n---\n请帮我处理一下这段文字。"

user_input = st.chat_input("说点什么，或者把左边识别的文字粘贴进来...", key="chat_input")

# 如果用户直接在输入框里点了发送（即使内容是空的，但如果有ocr结果也要处理）
# 这里稍微简化处理，依赖用户手动复制或输入，更稳定
if user_input:
    if not api_key:
        st.error("🚫 密钥未就位。")
    else:
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.write(user_input)

        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
        
        system_prompt = """
        你现在的身份是【程凯】。你极其聪明、善良且博学，说话风格抽象幽默，自带一种高级的优雅感。
        注意：绝对不要提物理。如果用户发给你一段识别出来的文字，请结合你的人设进行有趣或有深度的回复。
        """

        with st.chat_message("assistant"):
            with st.spinner("Kai is thinking..."):
                try:
                    response = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[
                            {"role": "system", "content": system_prompt},
                            *st.session_state.messages 
                        ]
                    )
                    result = response.choices[0].message.content
                    st.write(result)
                    st.session_state.messages.append({"role": "assistant", "content": result})
                    speak_text(result)
                    
                except Exception as e:
                    st.error(f"连接波动：{e}")